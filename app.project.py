import streamlit as st
import pytesseract
from PIL import Image
from docx import Document
import tempfile
import re

from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader

# -------- Tesseract Path --------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# -------- Page Config --------
st.set_page_config(page_title="FusionRAG Doc Chat", layout="wide")

# -------- ULTRA ANIMATED CSS --------
st.markdown("""
<style>

/* Animated moving background */
.stApp {
 background: linear-gradient(-45deg,#0f0c29,#240046,#3c096c,#000000,#10002b);
 background-size: 500% 500%;
 animation: gradientMove 15s ease infinite;
 color:white;
}

@keyframes gradientMove{
 0%{background-position:0% 50%}
 50%{background-position:100% 50%}
 100%{background-position:0% 50%}
}

/* Neon Title */
h1{
 text-align:center;
 font-size:50px;
 color: #FFFFFF;
 text-shadow: 0 0 10px #FFFFFF, 0 0 20px #FFFFFF;
 animation:pulse 2s infinite alternate;
}

@keyframes pulse{
 from{ text-shadow:0 0 15px #9d4edd;}
 to{ text-shadow:0 0 35px #e0aaff;}
}

/* Hover glow headings */
h2,h3{
 transition:0.3s;
 cursor:pointer;
}
h2:hover,h3:hover{
 color:#e0aaff;
 text-shadow:0 0 20px #c77dff;
 transform:scale(1.05);
}

/* Input neon */
.stTextInput input{
 background:#111 !important;
 color:white !important;
 border:1px solid #9d4edd !important;
 box-shadow:0 0 15px #9d4edd;
 border-radius:12px;
}

/* Sidebar */
section[data-testid="stSidebar"]{
 background:#0b0618;
}

/* Answer Cards */
.answer-card{
 padding:25px;
 border-radius:20px;
 margin:20px 0;
 background:rgba(20,10,50,0.9);
 border:1px solid #7b2cbf;
 box-shadow:0 0 25px #7b2cbf;
 transition:0.4s;
 animation:fadeIn 0.6s ease-in;
}

.answer-card:hover{
 transform:scale(1.03);
 box-shadow:0 0 50px #c77dff;
}

/* Fade animation */
@keyframes fadeIn{
 from{opacity:0; transform:translateY(20px);}
 to{opacity:1; transform:translateY(0);}
}

/* Highlighted words */
.highlight{
 background:linear-gradient(90deg,#ff00ff,#9d4edd);
 padding:4px 8px;
 border-radius:6px;
 font-weight:bold;
}
/* TextInput label */
.stTextInput label {
  color: #e0aaff !important;
  font-size: 20px !important;
  font-weight: bold;
  text-shadow: 0 0 10px #9d4edd;
}

/* Slider label */
.stSlider label {
  color: #e0aaff !important;
  font-size: 18px !important;
  font-weight: bold;
  text-shadow: 0 0 10px #9d4edd;
}

/* Slider value number */
.stSlider span {
  color: #ffffff !important;
  font-weight: bold;
}

/* Optional: subheader (Top-K Answers output title) */
.stSubheader {
  color: #e0aaff !important;
  text-shadow: 0 0 15px #c77dff;
}
/* TextInput label */
.stTextInput label {
  color: #e0aaff !important;
  font-size: 20px !important;
  font-weight: bold;
  text-shadow: 0 0 10px #9d4edd;
}

/* Slider label */
.stSlider label {
  color: #e0aaff !important;
  font-size: 18px !important;
  font-weight: bold;
  text-shadow: 0 0 10px #9d4edd;
}

/* Sidebar "Upload Files" header text */
section[data-testid="stSidebar"] h2 {
  color: #ffffff !important;
}

/* Sidebar "Upload PDFs, DOCX, Images" text */
section[data-testid="stSidebar"] label {
    color: #ffffff !important;
}



</style>
""", unsafe_allow_html=True)

st.title("Fusion RAG Doc Chat")

# -------- Sidebar --------
st.sidebar.header("📂 Upload Files")

files = st.sidebar.file_uploader(
    "Upload PDFs, DOCX, Images",
    type=["pdf","docx","png","jpg","jpeg"],
    accept_multiple_files=True
)

top_k = st.sidebar.slider("Top-K Answers",1,10,3)

# -------- Extractors --------
def extract_docx(file):
    doc = Document(file)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_image(file):
    img = Image.open(file)
    return pytesseract.image_to_string(img)

def load_files(files):
    docs=[]
    for file in files:
        ext=file.name.split(".")[-1].lower()

        if ext=="pdf":
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                tmp.write(file.read())
                loader=PyPDFLoader(tmp.name)
                docs.extend([d for d in loader.load() if d.page_content.strip()])

        elif ext=="docx":
            text=extract_docx(file)
            if text.strip():
                docs.append(LCDocument(page_content=text))

        elif ext in ["png","jpg","jpeg"]:
            text=extract_image(file)
            if text.strip():
                docs.append(LCDocument(page_content=text))
    return docs

# -------- Highlight Function --------
def highlight(text,query):
    for word in query.split():
        pattern=re.compile(re.escape(word),re.IGNORECASE)
        text=pattern.sub(f'<span class="highlight">{word}</span>',text)
    return text

# -------- Session DB --------
if "db" not in st.session_state:
    st.session_state.db=None

# -------- Build DB --------
if files and st.session_state.db is None:
    with st.spinner("⚡ Processing Documents..."):
        raw_docs=load_files(files)

        if not raw_docs:
            st.error("No readable text found.")
            st.stop()

        splitter = RecursiveCharacterTextSplitter(
    chunk_size=1200,
    chunk_overlap=250,
    separators=["\n\n", "\n", ". ", "? ", "! "]
    )

        chunks=splitter.split_documents(raw_docs)

        embeddings=HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )

        st.session_state.db=FAISS.from_documents(chunks,embeddings)

    st.success("✅ Documents Indexed!")

# -------- Query --------
query=st.text_input("💬 Ask from your documents")

if query and st.session_state.db:

    # get more chunks than needed
    results = st.session_state.db.similarity_search(query, k=top_k * 2)

    st.subheader(f"✨ Top {top_k} Answers")

    for i in range(top_k):

        # merge 2 chunks per answer
        combined = results[i].page_content

        if i + top_k < len(results):
            combined += " " + results[i + top_k].page_content

        # keep full sentences only
        sentences = re.split(r'(?<=[.!?]) +', combined)
        answer = " ".join(sentences[:5])  # first 5 full sentences

        # highlight
        for word in query.split():
            answer = re.sub(
                f"(?i)({word})",
                r'<span class="highlight">\1</span>',
                answer
            )

        st.markdown(f"""
        <div class="answer-card">
        <h3>Answer {i+1}</h3>
        <p style="font-size:18px; line-height:1.8; text-align:justify;">
        {answer}
        </p>
        </div>
        """, unsafe_allow_html=True)gi