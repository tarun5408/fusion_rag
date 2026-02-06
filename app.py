import streamlit as st
import pytesseract
from PIL import Image
from docx import Document
import tempfile

from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader

# ---------- Tesseract ----------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ---------- Page ----------
st.set_page_config(page_title="FusionRAG AI", layout="wide")

# ---------- CSS UI ----------
st.markdown("""
<style>
.stApp {
 background: linear-gradient(135deg,#0f0c29,#302b63,#24243e);
 color:white;
}

.answer-card{
 padding:20px;
 border-radius:15px;
 margin:15px 0;
 background: rgba(255,255,255,0.08);
 backdrop-filter: blur(12px);
 box-shadow: 0 0 25px #9d4edd;
}
</style>
""", unsafe_allow_html=True)

st.title("✨ FusionRAG Multi-Doc AI")

# ---------- Sidebar ----------
st.sidebar.header("📂 Upload Files")

files = st.sidebar.file_uploader(
    "Upload PDFs, DOCX, Images",
    type=["pdf","docx","png","jpg","jpeg"],
    accept_multiple_files=True
)

top_k = st.sidebar.slider("Top-K Answers",1,10,3)

# ---------- Extractors ----------
def extract_docx(file):
    doc = Document(file)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_image(file):
    img = Image.open(file)
    return pytesseract.image_to_string(img)

def load_files(files):
    docs=[]
    for file in files:
        ext=file.name.split(".")[-1].lower()

        if ext=="pdf":
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                tmp.write(file.read())
                loader=PyPDFLoader(tmp.name)
                pdf_docs=loader.load()
                docs.extend([d for d in pdf_docs if d.page_content.strip()])

        elif ext=="docx":
            text=extract_docx(file)
            if text.strip():
                docs.append(LCDocument(page_content=text))

        elif ext in ["png","jpg","jpeg"]:
            text=extract_image(file)
            if text.strip():
                docs.append(LCDocument(page_content=text))
    return docs

# ---------- Build DB once ----------
if "db" not in st.session_state:
    st.session_state.db=None

if files and st.session_state.db is None:

    with st.spinner("🔄 Processing files..."):
        raw_docs=load_files(files)

        if not raw_docs:
            st.error("No text found.")
            st.stop()

        splitter=RecursiveCharacterTextSplitter(
            chunk_size=700,
            chunk_overlap=150
        )

        chunks=splitter.split_documents(raw_docs)

        embeddings=HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )

        st.session_state.db=FAISS.from_documents(chunks,embeddings)

    st.success("✅ Files processed!")

# ---------- Query ----------
query=st.text_input("💬 Ask question from your documents")

if query and st.session_state.db:

    results=st.session_state.db.similarity_search(query,k=top_k)

    st.subheader(f"📌 Top {top_k} Answers")

    if not results:
        st.warning("No matches found.")
    else:
        for i,res in enumerate(results,1):
            st.markdown(f"""
            <div class="answer-card">
            <h4>Answer {i}</h4>
            {res.page_content[:600]}
            </div>
            """,unsafe_allow_html=True)

if not files:
    st.info("👈 Upload documents to start")